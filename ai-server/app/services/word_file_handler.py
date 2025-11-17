"""
Word 파일 생성 및 AI 서버 URL 생성 서비스
Word 파일을 생성하고 AI 서버에서 직접 서빙할 수 있는 URL을 반환합니다.
"""

import os
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import uuid
from dotenv import load_dotenv
from urllib.parse import quote

# .env 파일 로드
load_dotenv(verbose=True)


def create_word_file(cover_letter_text: str, cover_letter_data: Dict[str, Any], base_dir: Path = None) -> Dict[str, Any]:
    """
    자기소개서 텍스트를 Word 파일로 생성합니다.
    
    Args:
        cover_letter_text: 자기소개서 텍스트 내용
        cover_letter_data: 자기소개서 메타데이터
        base_dir: 파일을 저장할 기본 디렉토리 (기본값: files/resumes)
    
    Returns:
        파일 메타데이터 (filename, filepath 등)
    """
    try:
        # 파일 저장 디렉토리 설정
        if base_dir is None:
            base_dir = Path("files") / "resumes"
        base_dir.mkdir(parents=True, exist_ok=True)
        
        # 파일명 생성
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        position = cover_letter_data.get("position", "자기소개서")
        # 파일명에 사용할 수 없는 문자 제거
        position_clean = "".join(c for c in position if c.isalnum() or c in (' ', '-', '_')).strip()
        filename = f"자기소개서_{position_clean}_{timestamp}.docx"
        filepath = base_dir / filename
        
        # Word 문서 생성
        doc = Document()
        
        # 기본 스타일 설정
        style = doc.styles['Normal']
        style.font.name = '맑은 고딕'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        
        # 제목 추가
        heading = doc.add_heading('자기소개서', 0)
        for run in heading.runs:
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        
        # 본문 추가
        paragraphs = cover_letter_text.split('\n')
        for para in paragraphs:
            p = doc.add_paragraph()
            run = p.add_run(para)  # strip() 제거! 공백 유지
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        
        # 문서 저장
        doc.save(str(filepath))
        
        # 메타데이터 생성
        cover_letter_id = str(uuid.uuid4())
        metadata = {
            "id": cover_letter_id,
            "filename": filename,
            "filepath": str(filepath.absolute()),
            "created_at": datetime.now().isoformat(),
            "data": cover_letter_data,
            "status": "completed"
        }
        
        print(f"✅ Word 파일 생성 완료: {filepath}")
        return metadata
        
    except Exception as e:
        print(f"❌ Word 파일 생성 오류: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "error": str(e),
            "status": "error"
        }


def create_ai_server_file_url(filename: str, server_url: str = None) -> str:
    """
    AI 서버에서 파일에 접근할 수 있는 URL을 생성합니다.
    
    Args:
        filename: 파일명
        server_url: AI 서버의 기본 URL (기본값: .env에서 가져오거나 localhost:8000)
    
    Returns:
        AI 서버 파일 URL
    """
    if server_url is None:
        server_url = os.getenv("AI_SERVER_URL", "http://localhost:8000")
    
    # 파일명 URL 인코딩 (한글 등 특수문자 처리)
    encoded_filename = quote(filename, safe='')
    
    # URL 생성
    file_url = f"{server_url}/files/resumes/{encoded_filename}"
    
    return file_url


def create_word_file_and_url(cover_letter_text: str, cover_letter_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Word 파일을 생성하고 AI 서버에서 접근할 수 있는 URL을 반환합니다.
    
    Args:
        cover_letter_text: 자기소개서 텍스트 내용
        cover_letter_data: 자기소개서 메타데이터
    
    Returns:
        {
            "url": AI 서버 파일 URL,
            "filename": 파일명,
            "filepath": 파일 경로,
            "status": "completed" or "error"
        }
    """
    try:
        # 파일 저장 디렉토리
        base_dir = Path("files") / "resumes"
        
        # 1. Word 파일 생성
        word_metadata = create_word_file(cover_letter_text, cover_letter_data, base_dir)
        
        if word_metadata.get("status") == "error":
            return {
                "url": None,
                "filename": None,
                "filepath": None,
                "status": "error",
                "error": word_metadata.get("error")
            }
        
        file_path = word_metadata.get("filepath")
        filename = word_metadata.get("filename")
        
        if not file_path or not filename:
            return {
                "url": None,
                "filename": filename,
                "filepath": None,
                "status": "error",
                "error": "파일 경로 또는 파일명이 생성되지 않았습니다."
            }
        
        # 파일 경로가 상대 경로인 경우 절대 경로로 변환
        if not os.path.isabs(file_path):
            file_path = os.path.abspath(file_path)
        
        # 2. AI 서버 파일 URL 생성
        file_url = create_ai_server_file_url(filename)
        
        print(f"✅ Word 파일 URL 생성 완료: {file_url}")
        
        return {
            "url": file_url,
            "filename": filename,
            "filepath": file_path,
            "status": "completed"
        }
        
    except Exception as e:
        print(f"❌ Word 파일 생성 및 URL 변환 오류: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "url": None,
            "filename": None,
            "filepath": None,
            "status": "error",
            "error": str(e)
        }
