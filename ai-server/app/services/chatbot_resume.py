##url로 돌려주기 !!!
## 자기소개서 작성 챗봇 AI
## 넥스터(자기소개서 챗봇)를 통해 사용자와 대화하며 자기소개서를 작성합니다.

import os
import json
import uuid
from datetime import datetime
from typing import Dict, Any, Optional, List
from pathlib import Path
from dotenv import load_dotenv
from openai import OpenAI
from fpdf import FPDF
from docx import Document

# .env 파일 로드
load_dotenv(verbose=True)

# OpenAI API 키 확인
openai_api_key = os.getenv("OPENAI_API_KEY")
if not openai_api_key:
    raise ValueError("OPENAI_API_KEY가 .env 파일에 설정되지 않았습니다.")

client = OpenAI(api_key=openai_api_key)

# 자기소개서 필드 정의
COVER_LETTER_FIELDS = {
    "position": "직무 목표",
    "skills": "기술 스택",
    "experience": "최근 경력",
    "achievements": "주요 성과",
    "motivation": "지원 동기",
    "strengths": "강점",
    "personality": "성격/특징",
    "future_plans": "향후 계획"
}

# 완료를 나타내는 키워드
COMPLETION_KEYWORDS = ["완료", "저장", "끝", "종료", "done", "save", "finish", "complete"]

# 확인을 나타내는 키워드
CONFIRMATION_KEYWORDS = ["맞아", "네", "예", "ok", "okay", "좋아", "맞아요", "네요", "예요", "그래", "그래요", "확인", "yes", "y"]

# 자기소개서 작성 의도 확인 키워드
COVER_LETTER_INTENT_KEYWORDS = ["자기소개서", "자소서", "지원서", "cover letter", "이력서", "resume"]

def is_cover_letter_intent(user_message: str) -> bool:
    """사용자 메시지가 자기소개서 작성 의도인지 확인합니다."""
    if not user_message:
        return False
    message_lower = user_message.strip().lower()
    return any(keyword in message_lower for keyword in COVER_LETTER_INTENT_KEYWORDS)

def is_completion_request(user_message: str) -> bool:
    """사용자 메시지가 완료 요청인지 확인합니다."""
    if not user_message:
        return False
    message_lower = user_message.strip().lower()
    return any(keyword in message_lower for keyword in COMPLETION_KEYWORDS)

def is_confirmation(user_message: str) -> bool:
    """사용자 메시지가 확인 응답인지 확인합니다."""
    if not user_message:
        return False
    message_lower = user_message.strip().lower()
    return any(keyword in message_lower for keyword in CONFIRMATION_KEYWORDS)

def detect_user_intent_with_llm(user_message: str, conversation_context: Optional[str] = None) -> Dict[str, Any]:
    """LLM을 사용하여 사용자의 의도를 파악합니다."""
    try:
        context = ""
        if conversation_context:
            context = f"\n대화 맥락: {conversation_context}"
        
        prompt = f"""당신은 자기소개서 작성을 도와주는 AI 챗봇 '넥스터'입니다.

사용자가 "안녕하세요! 저는 자기소개서 작성을 도와주는 넥스터입니다. 자기소개서 작성을 원하시나요?"라고 물어본 후, 사용자가 다음과 같이 답했습니다:

사용자 메시지: {user_message}
{context}

사용자의 답변을 분석하여 자기소개서 작성을 원하는지 확인해주세요.
- "응", "네", "예", "좋아", "그래", "응응" 등의 긍정 응답은 자기소개서 작성을 원하는 것으로 간주합니다.
- "아니", "안 해", "싫어" 등의 부정 응답은 원하지 않는 것으로 간주합니다.
- 명확하지 않은 경우도 긍정적으로 해석합니다.

JSON 형식으로 응답하세요:
{{
  "wants_cover_letter": true/false,
  "confidence": "high"/"medium"/"low",
  "reasoning": "판단 근거"
}}"""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.3
        )
        
        result = json.loads(response.choices[0].message.content)
        return result
        
    except Exception as e:
        print(f"의도 파악 오류: {str(e)}")
        # 오류 발생 시 기본적으로 확인 응답으로 간주
        return {
            "wants_cover_letter": True,
            "confidence": "low",
            "reasoning": "오류로 인한 기본값"
        }

def format_basic_info_summary(cover_letter_data: Dict[str, Any]) -> str:
    """기본 정보를 요약 형식으로 포맷팅합니다."""
    summary_parts = []
    
    position = cover_letter_data.get("position")
    if position:
        summary_parts.append(f"• 직무 목표: {position}")
    
    skills = cover_letter_data.get("skills", [])
    if skills:
        skills_str = ", ".join(skills) if isinstance(skills, list) else skills
        summary_parts.append(f"• 기술 스택: {skills_str}")
    
    experience = cover_letter_data.get("experience")
    if experience:
        summary_parts.append(f"• 최근 경력: {experience}")
    
    achievements = cover_letter_data.get("achievements", [])
    if achievements:
        achievements_str = ", ".join(achievements) if isinstance(achievements, list) else achievements
        summary_parts.append(f"• 주요 성과: {achievements_str}")
    
    return "\n".join(summary_parts) if summary_parts else "아직 입력된 정보가 없습니다."

def convert_metadata_to_cover_letter_data(metadata: Dict[str, Any]) -> Dict[str, Any]:
    """file_analysis.py의 메타데이터를 자기소개서 데이터 형식으로 변환합니다."""
    project = metadata.get("project", {})
    
    cover_letter_data = {
        "position": None,
        "skills": [],
        "experience": None,
        "achievements": [],
        "motivation": None,
        "strengths": [],
        "personality": None,
        "future_plans": None
    }
    
    # 메타데이터 매핑
    # tools → skills
    if project.get("tools"):
        cover_letter_data["skills"] = project["tools"] if isinstance(project["tools"], list) else [project["tools"]]
    
    # roles → position 또는 experience
    if project.get("roles"):
        roles = project["roles"] if isinstance(project["roles"], list) else [project["roles"]]
        if roles:
            # 첫 번째 역할을 직무로
            cover_letter_data["position"] = roles[0] if len(roles) > 0 else None
            # 나머지 역할을 경력으로
            if len(roles) > 1:
                cover_letter_data["experience"] = ", ".join(roles[1:])
    
    # category → position (roles가 없을 경우)
    if not cover_letter_data["position"] and project.get("category"):
        cover_letter_data["position"] = project["category"]
    
    # achievements → achievements
    if project.get("achievements"):
        cover_letter_data["achievements"] = project["achievements"] if isinstance(project["achievements"], list) else [project["achievements"]]
    
    # description → experience 또는 motivation
    if project.get("description"):
        if not cover_letter_data["experience"]:
            cover_letter_data["experience"] = project["description"]
        else:
            cover_letter_data["motivation"] = project["description"]
    
    # title → experience에 추가 정보로 활용 가능
    if project.get("title") and not cover_letter_data["experience"]:
        cover_letter_data["experience"] = project["title"]
    
    return cover_letter_data

def create_sample_metadata() -> Dict[str, Any]:
    """테스트용 임의의 메타데이터를 생성합니다."""
    return {
        "project": {
            "title": "AI 기반 마케팅 캠페인 프로젝트",
            "category": "마케팅/기획",
            "tags": ["마케팅", "AI", "데이터 분석"],
            "roles": ["마케팅 기획자", "프로젝트 매니저"],
            "achievements": [
                "지역 마케팅 캠페인 참여율 200% 향상",
                "SNS 콘텐츠 전략 프로젝트로 광고 예산 효율성 개선"
            ],
            "tools": ["Google Analytics", "Figma", "Notion", "Photoshop"],
            "description": "콘텐츠 캠페인 기획 및 브랜드 홍보 전략 수립을 담당했습니다. 데이터 기반 의사결정을 통해 마케팅 성과를 극대화했습니다."
        },
        "status": "analyzed"
    }

def extract_and_update_cover_letter_data_with_llm(
    cover_letter_data: Dict[str, Any],
    user_message: str,
    conversation_history: List[Dict[str, str]]
) -> Dict[str, Any]:
    """LLM을 사용하여 사용자 메시지에서 자기소개서 정보를 추출하고 업데이트합니다."""
    try:
        current_data_str = json.dumps(cover_letter_data, ensure_ascii=False, indent=2)
        
        # 대화 히스토리 요약
        history_summary = "\n".join([
            f"{msg['role']}: {msg['content']}" 
            for msg in conversation_history[-5:]
        ])
        
        prompt = f"""당신은 자기소개서 작성을 도와주는 AI 챗봇 '넥스터'입니다.

현재 수집된 정보:
{current_data_str}

최근 대화:
{history_summary}

사용자 메시지: {user_message}

사용자의 메시지를 분석하여 다음 정보를 추출하고 업데이트하세요:
- position: 직무 목표 (예: 마케팅/기획, 개발자, 디자이너 등)
- skills: 기술 스택 (배열, 예: ["Google Analytics", "Figma", "Notion"])
- experience: 최근 경력/경험
- achievements: 주요 성과 (배열)
- motivation: 지원 동기
- strengths: 강점 (배열)
- personality: 성격/특징
- future_plans: 향후 계획

사용자가 불완전한 정보를 제공했거나 기억이 안 난다고 하면, 대화를 통해 자연스럽게 추가 정보를 물어보세요.

JSON 형식으로 응답하세요:
{{
  "updated_data": {{
    "position": "직무 또는 null",
    "skills": ["기술1", "기술2"] 또는 [],
    "experience": "경력/경험 또는 null",
    "achievements": ["성과1", "성과2"] 또는 [],
    "motivation": "지원 동기 또는 null",
    "strengths": ["강점1", "강점2"] 또는 [],
    "personality": "성격/특징 또는 null",
    "future_plans": "향후 계획 또는 null"
  }},
  "response_message": "사용자에게 자연스럽게 대화를 이어갈 수 있는 메시지",
  "needs_more_info": true/false
}}

기존 값이 있으면 유지하되, 새로운 정보가 제공되면 업데이트하세요."""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.7
        )
        
        result = json.loads(response.choices[0].message.content)
        
        # 데이터 업데이트 (null이 아닌 값만 업데이트)
        updated_data = json.loads(json.dumps(cover_letter_data))
        
        llm_data = result.get("updated_data", {})
        
        for field, value in llm_data.items():
            if value is not None:
                if isinstance(value, list) and len(value) > 0:
                    # 배열 필드: 기존 값과 병합 (중복 제거)
                    existing = set(updated_data.get(field, []))
                    new_values = set(value)
                    updated_data[field] = list(existing | new_values)
                elif not isinstance(value, list) and value != "":
                    # 단일 값 필드
                    updated_data[field] = value
        
        return {
            "updated_data": updated_data,
            "response_message": result.get("response_message", ""),
            "needs_more_info": result.get("needs_more_info", False)
        }
        
    except Exception as e:
        print(f"자기소개서 데이터 추출 오류: {str(e)}")
        return {
            "updated_data": cover_letter_data,
            "response_message": "죄송합니다. 이해하지 못했습니다. 다시 말씀해주실 수 있을까요?",
            "needs_more_info": True
        }

def generate_cover_letter_draft(cover_letter_data: Dict[str, Any], writing_style: str = "자연스럽고 전문적인") -> str:
    """수집된 정보를 바탕으로 자기소개서 초안을 생성합니다."""
    try:
        data_str = json.dumps(cover_letter_data, ensure_ascii=False, indent=2)
        
        # 프로젝트가 여러 개인지 확인
        projects = cover_letter_data.get("projects", [])
        project_instruction = ""
        
        if len(projects) > 1:
            project_instruction = "\n\n**중요: 데이터에 포함된 프로젝트들을 각각 구분해서 자기소개서에 반영하세요. 프로젝트가 여러 개인 경우, 각 프로젝트를 별도 문단으로 작성하거나 구분해서 설명해주세요.**"
        elif len(projects) == 1:
            project_instruction = "\n\n**중요: 데이터에 포함된 프로젝트 정보를 활용하여 자기소개서를 작성하세요.**"
        
        prompt = f"""다음 정보를 바탕으로 {writing_style} 문체로 자기소개서 초안을 작성해주세요.

수집된 정보:
{data_str}
{project_instruction}

자기소개서는 다음을 포함해야 합니다:
1. 지원 동기 및 직무에 대한 관심
2. 보유 기술과 경험
3. 주요 성과와 결과
4. 강점과 특징
5. 향후 계획

문체는 {writing_style} 느낌으로 작성해주세요.
구체적이고 설득력 있게 작성하되, 자연스럽게 표현해주세요."""

        response = client.chat.completions.create(
            model="gpt-4o-mini",  # 속도 우선으로 변경
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        print(f"자기소개서 초안 생성 오류: {str(e)}")
        return "자기소개서 초안 생성 중 오류가 발생했습니다."

def modify_cover_letter(cover_letter_text: str, modification_request: str) -> str:
    """사용자의 수정 요청을 반영하여 자기소개서를 수정합니다."""
    try:
        prompt = f"""다음 자기소개서를 사용자의 요청에 맞게 수정해주세요.

현재 자기소개서:
{cover_letter_text}

사용자 요청: {modification_request}

**수정 가이드라인:**
1. 사용자가 "마지막에 ~로 끝내줘" 또는 "~로 끝내주세요"라고 하면:
   - 자기소개서의 마지막 문장을 사용자가 요청한 내용으로 정확히 교체하세요
   - "감사합니다."는 항상 맨 마지막에 유지하세요

2. 사용자가 "~를 추가해줘"라고 하면:
   - 적절한 위치에 해당 내용을 자연스럽게 추가하세요

3. 사용자가 "~를 바꿔줘" 또는 "~를 수정해줘"라고 하면:
   - 해당 부분을 찾아서 정확히 교체하세요

4. 전체적인 문체와 톤은 유지하되, 요청된 부분은 **반드시** 정확히 반영하세요.

수정된 자기소개서 전문만 출력하세요. 설명이나 추가 코멘트는 불필요합니다."""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        print(f"자기소개서 수정 오류: {str(e)}")
        return cover_letter_text

def save_cover_letter_as_pdf(cover_letter_text: str, cover_letter_data: Dict[str, Any]) -> Dict[str, Any]:
    """자기소개서를 PDF 파일로 저장합니다."""
    try:
        # 저장 디렉토리 생성
        output_dir = Path("output/cover_letters")
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # 고유 ID 생성
        cover_letter_id = str(uuid.uuid4())
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # 파일명 생성
        position = cover_letter_data.get("position", "자기소개서")
        # 파일명에 사용할 수 없는 문자 제거
        position_clean = "".join(c for c in position if c.isalnum() or c in (' ', '-', '_')).strip()
        filename = f"자기소개서_{position_clean}_{timestamp}.pdf"
        filepath = output_dir / filename
        
        # PDF 생성
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # 한글 폰트 설정 (한글 지원을 위해)
        try:
            # Windows의 경우 기본 폰트 사용
            pdf.set_font("Arial", size=12)
        except:
            pdf.set_font("Arial", size=12)
        
        # 텍스트를 줄 단위로 분리하여 추가
        lines = cover_letter_text.split('\n')
        for line in lines:
            if line.strip():
                # 한글 인코딩 처리
                try:
                    pdf.multi_cell(0, 8, line.encode('latin-1', 'replace').decode('latin-1'))
                except:
                    # 한글이 포함된 경우 다른 방식으로 처리
                    pdf.multi_cell(0, 8, line)
            else:
                pdf.ln(4)
        
        # PDF 저장
        pdf.output(str(filepath))
        
        # 메타데이터 생성
        metadata = {
            "id": cover_letter_id,
            "filename": filename,
            "filepath": str(filepath),
            "url": f"/cover_letters/{filename}",
            "created_at": datetime.now().isoformat(),
            "data": cover_letter_data,
            "status": "completed"
        }
        
        # 메타데이터 JSON 파일로 저장
        metadata_filepath = output_dir / f"{cover_letter_id}_metadata.json"
        with open(metadata_filepath, "w", encoding="utf-8") as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)
        
        return metadata
        
    except Exception as e:
        print(f"PDF 저장 오류: {str(e)}")
        return {
            "error": str(e),
            "status": "error"
        }

def save_cover_letter_as_word(cover_letter_text: str, cover_letter_data: Dict[str, Any]) -> Dict[str, Any]:
    """자기소개서를 Word 파일로 저장합니다."""
    try:
        # 현재 폴더에 저장
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # 파일명 생성
        position = cover_letter_data.get("position", "자기소개서")
        # 파일명에 사용할 수 없는 문자 제거
        position_clean = "".join(c for c in position if c.isalnum() or c in (' ', '-', '_')).strip()
        filename = f"자기소개서_{position_clean}_{timestamp}.docx"
        filepath = Path(filename)  # 현재 폴더에 저장
        
        # Word 문서 생성
        doc = Document()
        
        # 한글 폰트 설정
        from docx.shared import Pt
        from docx.oxml.ns import qn
        
        # 기본 스타일 설정
        style = doc.styles['Normal']
        style.font.name = '맑은 고딕'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        
        # 제목 추가
        heading = doc.add_heading('자기소개서', 0)
        for run in heading.runs:
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        
        # 본문 추가
        paragraphs = cover_letter_text.split('\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph()
                run = p.add_run(para.strip())
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            else:
                doc.add_paragraph()  # 빈 줄
        
        # 문서 저장
        doc.save(str(filepath))
        
        # 메타데이터 생성
        cover_letter_id = str(uuid.uuid4())
        metadata = {
            "id": cover_letter_id,
            "filename": filename,
            "filepath": str(filepath.absolute()),
            "url": f"/cover_letters/{filename}",
            "created_at": datetime.now().isoformat(),
            "data": cover_letter_data,
            "status": "completed"
        }
        
        return metadata
        
    except Exception as e:
        print(f"Word 파일 저장 오류: {str(e)}")
        return {
            "error": str(e),
            "status": "error"
        }

def process_cover_letter_chatbot(
    user_message: Optional[str] = None,
    cover_letter_data: Optional[Dict[str, Any]] = None,
    conversation_history: Optional[List[Dict[str, str]]] = None,
    current_state: str = "intent_confirmation",
    writing_style: Optional[str] = None,
    draft_cover_letter: Optional[str] = None,
    metadata: Optional[Dict[str, Any]] = None  # file_analysis.py의 메타데이터 추가
) -> Dict[str, Any]:
    """자기소개서 챗봇 메시지를 처리하고 응답을 반환합니다."""
    try:
        # 초기 데이터 설정
        if cover_letter_data is None:
            cover_letter_data = {
                "position": None,
                "skills": [],
                "experience": None,
                "achievements": [],
                "motivation": None,
                "strengths": [],
                "personality": None,
                "future_plans": None
            }
        
        # 메타데이터가 제공된 경우, 메타데이터를 기반으로 초기 데이터 설정
        if metadata and metadata.get("status") == "analyzed":
            metadata_based_data = convert_metadata_to_cover_letter_data(metadata)
            # 기존 데이터와 병합 (메타데이터 값이 있으면 우선 사용)
            for key, value in metadata_based_data.items():
                if value is not None and value != []:
                    if isinstance(value, list):
                        # 리스트인 경우 병합
                        existing = cover_letter_data.get(key, [])
                        if not isinstance(existing, list):
                            existing = []
                        cover_letter_data[key] = list(set(existing + value))
                    else:
                        # 단일 값인 경우 메타데이터 값 사용
                        if not cover_letter_data.get(key):
                            cover_letter_data[key] = value
        
        if conversation_history is None:
            conversation_history = []
        
        # 상태별 처리
        if current_state == "intent_confirmation":
            # 의도 확인 단계 - LLM을 사용하여 의도 파악
            if user_message:
                # LLM으로 의도 파악
                intent_result = detect_user_intent_with_llm(user_message)
                wants_cover_letter = intent_result.get("wants_cover_letter", False)
                
                if wants_cover_letter or is_cover_letter_intent(user_message):
                    # 메타데이터 기반 정보가 있으면 표시
                    basic_info = format_basic_info_summary(cover_letter_data)
                    if basic_info != "아직 입력된 정보가 없습니다.":
                        # 메타데이터 기반 정보가 충분한지 확인
                        has_position = cover_letter_data.get("position")
                        has_skills = cover_letter_data.get("skills") and len(cover_letter_data.get("skills", [])) > 0
                        has_experience = cover_letter_data.get("experience")
                        
                        # position이 있고, (skills 또는 experience) 중 하나만 있어도 바로 문체 선택으로
                        if has_position and (has_skills or has_experience):
                            message = f"확인된 기본 정보\n\n{basic_info}\n\n이 정보를 기반으로 자기소개서를 작성할게요.\n\n문체는 어떤 스타일로 원하시나요?"
                            return {
                                "message": message,
                                "updated_data": cover_letter_data,
                                "status": "style_selection",
                                "next_state": "style_selection"
                            }
                        else:
                            message = f"안녕하세요! 저는 자기소개서 작성을 도와주는 넥스터입니다.\n\n분석된 정보를 확인했습니다:\n\n{basic_info}\n\n추가로 필요한 정보를 빠르게 수집하겠습니다."
                    else:
                        message = "안녕하세요! 저는 자기소개서 작성을 도와주는 넥스터입니다.\n\n자기소개서 작성을 위해 몇 가지 정보가 필요합니다.\n\n먼저 지원하시는 직무 목표를 알려주세요. (예: 마케팅/기획, 개발자, 디자이너 등)"
                    
                    return {
                        "message": message,
                        "updated_data": cover_letter_data,
                        "status": "collecting_info",
                        "next_state": "collecting_info"
                    }
                else:
                    return {
                        "message": "안녕하세요! 저는 자기소개서 작성을 도와주는 넥스터입니다. 자기소개서 작성을 원하시나요?",
                        "updated_data": cover_letter_data,
                        "status": "intent_confirmation",
                        "next_state": "intent_confirmation"
                    }
            else:
                return {
                    "message": "안녕하세요! 저는 자기소개서 작성을 도와주는 넥스터입니다. 자기소개서 작성을 원하시나요?",
                    "updated_data": cover_letter_data,
                    "status": "intent_confirmation",
                    "next_state": "intent_confirmation"
                }
        
        elif current_state == "collecting_info":
            # 정보 수집 단계
            if user_message:
                result = extract_and_update_cover_letter_data_with_llm(
                    cover_letter_data,
                    user_message,
                    conversation_history
                )
                
                updated_data = result["updated_data"]
                response_message = result["response_message"]
                needs_more_info = result["needs_more_info"]
                
                # 필수 정보 확인 조건 완화 (position + skills 또는 experience 중 하나만 있어도 진행)
                has_position = updated_data.get("position")
                has_skills = updated_data.get("skills") and len(updated_data.get("skills", [])) > 0
                has_experience = updated_data.get("experience")
                
                # position이 있고, (skills 또는 experience) 중 하나만 있어도 진행
                if has_position and (has_skills or has_experience) and not needs_more_info:
                    # 기본 정보가 충분히 수집됨 - 기본 정보 확인 단계로
                    basic_info = format_basic_info_summary(updated_data)
                    return {
                        "message": f"확인된 기본 정보\n\n{basic_info}\n\n이 정보를 기반으로 자기소개서를 작성할게요.\n\n문체는 어떤 스타일로 원하시나요?",
                        "updated_data": updated_data,
                        "status": "style_selection",
                        "next_state": "style_selection"
                    }
                else:
                    # 추가 정보 필요
                    return {
                        "message": response_message,
                        "updated_data": updated_data,
                        "status": "collecting_info",
                        "next_state": "collecting_info"
                    }
            else:
                return {
                    "message": "지원하시는 직무 목표를 알려주세요.",
                    "updated_data": cover_letter_data,
                    "status": "collecting_info",
                    "next_state": "collecting_info"
                }
        
        elif current_state == "style_selection":
            # 문체 선택 단계
            if user_message:
                # 문체 저장하고 바로 초안 생성
                writing_style = user_message
                
                # 초안 즉시 생성
                draft = generate_cover_letter_draft(cover_letter_data, writing_style or "자연스럽고 전문적인")
                
                return {
                    "message": f"AI 초안 미리보기 ✅\n\n\"{draft}\"\n\n---\n\n어때요? 마음에 드시나요?\n수정하고 싶거나 추가하고 싶은 내용이 있으면 알려주세요!\n완성했다면 '완료' 또는 '저장'이라고 말씀해주세요.",
                    "updated_data": cover_letter_data,
                    "status": "draft_preview",
                    "next_state": "draft_revision",
                    "draft_cover_letter": draft,
                    "writing_style": writing_style
                }
            else:
                return {
                    "message": "문체는 어떤 스타일로 원하시나요? (예: 자연스럽고 전문적인, 격식 있는, 친근한 등)",
                    "updated_data": cover_letter_data,
                    "status": "style_selection",
                    "next_state": "style_selection"
                }
        
        elif current_state == "draft_preview":
            # 초안 미리보기 단계
            if not draft_cover_letter:
                # 초안 생성
                draft = generate_cover_letter_draft(cover_letter_data, writing_style or "자연스럽고 전문적인")
                return {
                    "message": f"AI 초안 미리보기 ✅\n\n\"{draft}\"\n\n---\n\n어때요? 마음에 드시나요?\n수정하고 싶거나 추가하고 싶은 내용이 있으면 알려주세요!\n완성했다면 '완료' 또는 '저장'이라고 말씀해주세요.",
                    "updated_data": cover_letter_data,
                    "status": "draft_preview",
                    "next_state": "draft_revision",
                    "draft_cover_letter": draft,
                    "writing_style": writing_style
                }
            else:
                # 이미 초안이 있음
                return {
                    "message": f"AI 초안 미리보기 ✅\n\n\"{draft_cover_letter}\"\n\n---\n\n어때요? 마음에 드시나요?\n수정하고 싶거나 추가하고 싶은 내용이 있으면 알려주세요!\n완성했다면 '완료' 또는 '저장'이라고 말씀해주세요.",
                    "updated_data": cover_letter_data,
                    "status": "draft_preview",
                    "next_state": "draft_revision",
                    "draft_cover_letter": draft_cover_letter,
                    "writing_style": writing_style
                }
        
        elif current_state == "draft_revision":
            # 초안 수정 단계
            if user_message:
                # 이전 AI 메시지 확인
                last_ai_message = ""
                if conversation_history:
                    for msg in reversed(conversation_history):
                        if msg.get("role") == "assistant":
                            last_ai_message = msg.get("content", "")
                            break
                
                # "이게 맞나요?" 메시지 이후 "네" 응답이면 바로 완료
                if ("이게 맞나요" in last_ai_message) and (is_confirmation(user_message) or "네" in user_message or "맞아요" in user_message):
                    # 바로 완료 상태로
                    return {
                        "message": "완료 ✅\n\nWord 파일을 생성 중입니다...",
                        "updated_data": cover_letter_data,
                        "status": "completed",
                        "next_state": "completed",
                        "draft_cover_letter": draft_cover_letter,
                        "writing_style": writing_style
                    }
                
                # "수정" 키워드가 있으면 수정 처리
                elif "수정" in user_message.lower() and ("이게 맞나요" in last_ai_message):
                    # 수정 요청
                    modified_draft = modify_cover_letter(draft_cover_letter, user_message)
                    return {
                        "message": f"수정 완료 ✅\n\n{modified_draft}\n\n이게 맞나요? 맞다면 네 라고 말해주시고 다시 수정을 원하면 수정이라고 말해주세요",
                        "updated_data": cover_letter_data,
                        "status": "draft_revision",
                        "next_state": "draft_revision",
                        "draft_cover_letter": modified_draft,
                        "writing_style": writing_style
                    }
                
                # 일반 확인 (초안에서 수정 없이 확인)
                elif is_confirmation(user_message) or "좋아" in user_message or "좋아요" in user_message:
                    # 수정 없이 확인 → final_confirmation
                    return {
                        "message": "최종 자기소개서를 확인하시겠어요?",
                        "updated_data": cover_letter_data,
                        "status": "final_confirmation",
                        "next_state": "final_confirmation",
                        "draft_cover_letter": draft_cover_letter,
                        "writing_style": writing_style
                    }
                
                else:
                    # 수정 요청
                    modified_draft = modify_cover_letter(draft_cover_letter, user_message)
                    return {
                        "message": f"수정 완료 ✅\n\n{modified_draft}\n\n이게 맞나요? 맞다면 네 라고 말해주시고 다시 수정을 원하면 수정이라고 말해주세요",
                        "updated_data": cover_letter_data,
                        "status": "draft_revision",
                        "next_state": "draft_revision",
                        "draft_cover_letter": modified_draft,
                        "writing_style": writing_style
                    }
            else:
                return {
                    "message": "수정이 필요하시면 알려주세요. 그렇지 않으면 '좋아' 또는 '확인'이라고 말씀해주세요.",
                    "updated_data": cover_letter_data,
                    "status": "draft_revision",
                    "next_state": "draft_revision",
                    "draft_cover_letter": draft_cover_letter,
                    "writing_style": writing_style
                }
        
        elif current_state == "final_confirmation":
            # 최종 확인 단계
            if user_message and (is_confirmation(user_message) or "pdf" in user_message.lower() or "다운로드" in user_message or "word" in user_message.lower() or "docx" in user_message.lower() or "저장" in user_message):
                # 파일 생성은 main.py에서 처리하므로 여기서는 메시지만 반환
                return {
                    "message": "완료 ✅\n\nWord 파일을 생성 중입니다...",
                    "updated_data": cover_letter_data,
                    "status": "completed",
                    "next_state": "completed",
                    "draft_cover_letter": draft_cover_letter
                }
            else:
                return {
                    "message": "최종 자기소개서를 확인하시겠어요? (Word 파일로 저장하려면 '예' 또는 '저장'이라고 말씀해주세요)",
                    "updated_data": cover_letter_data,
                    "status": "final_confirmation",
                    "next_state": "final_confirmation",
                    "draft_cover_letter": draft_cover_letter,
                    "writing_style": writing_style
                }
        
        elif current_state == "completed":
            # 완료 후 추가 질의응답 처리
            if user_message:
                # 간단한 질의응답 처리
                try:
                    prompt = f"""사용자가 자기소개서 작성을 완료했습니다.

현재 생성된 자기소개서:
{draft_cover_letter[:500] if draft_cover_letter else "없음"}...

사용자의 추가 질문: {user_message}

사용자의 질문에 친절하게 답변해주세요.
- 자기소개서를 다시 수정하고 싶다면: "새로운 자기소개서를 작성하려면 처음부터 다시 시작해주세요"라고 안내
- 파일 다운로드 관련 질문: "이미 생성된 파일을 다운로드하실 수 있습니다"라고 안내
- 일반적인 질문이면 친절하게 답변
- 자기소개서 작성 팁이나 조언을 요청하면 구체적으로 답변"""

                    response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.7,
                        max_tokens=500
                    )
                    
                    return {
                        "message": response.choices[0].message.content,
                        "updated_data": cover_letter_data,
                        "status": "completed",
                        "next_state": "completed",
                        "draft_cover_letter": draft_cover_letter
                    }
                    
                except Exception as e:
                    print(f"완료 후 질의응답 오류: {str(e)}")
                    return {
                        "message": "질문에 답변하는 중 오류가 발생했습니다. 다시 질문해주세요.",
                        "updated_data": cover_letter_data,
                        "status": "completed",
                        "next_state": "completed",
                        "draft_cover_letter": draft_cover_letter
                    }
            else:
                return {
                    "message": "자기소개서가 완성되었습니다! 추가로 궁금한 점이 있으시면 언제든 물어보세요.",
                    "updated_data": cover_letter_data,
                    "status": "completed",
                    "next_state": "completed",
                    "draft_cover_letter": draft_cover_letter
                }
        
        else:
            return {
                "message": "처리 중 오류가 발생했습니다.",
                "updated_data": cover_letter_data,
                "status": "error",
                "next_state": "intent_confirmation"
            }
        
    except Exception as e:
        print(f"챗봇 처리 오류: {str(e)}")
        return {
            "message": "죄송합니다. 처리 중 오류가 발생했습니다.",
            "updated_data": cover_letter_data if cover_letter_data else {},
            "status": "error",
            "next_state": "intent_confirmation"
        }

def main():
    """메인 함수 - 테스트용"""
    import sys
    
    # Windows 콘솔 인코딩 설정
    if sys.platform == 'win32':
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')
    
    print("=" * 60)
    print("넥스터 - 자기소개서 작성 챗봇")
    print("=" * 60)
    print()
    
    # 임의의 메타데이터 생성 (테스트용)
    metadata = create_sample_metadata()
    print("테스트용 메타데이터를 생성했습니다:")
    print(json.dumps(metadata, ensure_ascii=False, indent=2))
    print()
    
    # 초기 상태
    cover_letter_data = {
        "position": None,
        "skills": [],
        "experience": None,
        "achievements": [],
        "motivation": None,
        "strengths": [],
        "personality": None,
        "future_plans": None
    }
    
    conversation_history = []
    current_state = "intent_confirmation"
    writing_style = None
    draft_cover_letter = None
    
    # 맨 처음 넥스터 인사 메시지 출력
    initial_result = process_cover_letter_chatbot(
        None,  # user_message가 None이면 초기 인사
        cover_letter_data,
        conversation_history,
        current_state,
        writing_style,
        draft_cover_letter,
        metadata  # 메타데이터 전달
    )
    
    if initial_result.get("message"):
        print(f"Nexter: {initial_result['message']}\n")
        conversation_history.append({"role": "assistant", "content": initial_result['message']})
    
    while True:
        # draft_preview 상태이고 초안이 없으면 사용자 입력 없이 자동으로 초안 생성
        if current_state == "draft_preview" and not draft_cover_letter:
            user_input = None  # 사용자 입력 없이 처리
        else:
            user_input = input("사용자: ").strip()
            if not user_input:
                continue
        
        result = process_cover_letter_chatbot(
            user_input,
            cover_letter_data,
            conversation_history,
            current_state,
            writing_style,
            draft_cover_letter,
            metadata  # 메타데이터 전달
        )
        
        # 상태 업데이트
        current_state = result.get("next_state", current_state)
        cover_letter_data = result.get("updated_data", cover_letter_data)
        writing_style = result.get("writing_style", writing_style)
        draft_cover_letter = result.get("draft_cover_letter", draft_cover_letter)
        
        # 대화 히스토리 업데이트
        if user_input:
            conversation_history.append({"role": "user", "content": user_input})
        if result.get("message"):
            print(f"\nNexter: {result['message']}\n")
            conversation_history.append({"role": "assistant", "content": result['message']})
        
        # 완료 상태 확인
        if result.get("status") == "completed":
            print("\n" + "=" * 60)
            if result.get("metadata"):
                print("메타데이터:")
                print(json.dumps(result["metadata"], ensure_ascii=False, indent=2))
            break

if __name__ == "__main__":
    main()

